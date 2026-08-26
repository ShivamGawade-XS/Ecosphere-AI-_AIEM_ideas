from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path("/home/ubuntu/ecosphere-ai-website/EcoSphere_AI_AIEM_IDEAS_2026_Dossier.docx")


GREEN = "214D40"
MOSS = "6C8A72"
PAPER = "F7F5EE"
INK = "1F2B25"
GOLD = "A7791C"


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_rule(paragraph, color: str = GOLD) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run_font(run, size: float | None = None, bold: bool = False, color: str = INK, name: str = "Aptos") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)


def add_title(document: Document, title: str, subtitle: str) -> None:
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(50)
    title_p.paragraph_format.space_after = Pt(8)
    title_run = title_p.add_run(title)
    set_run_font(title_run, 30, True, GREEN, "Georgia")

    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run(subtitle)
    set_run_font(sub_run, 12, True, MOSS)
    add_rule(sub_p)

    tag_p = document.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_after = Pt(32)
    tag_run = tag_p.add_run("AIEM IDEAS 2026  •  REGISTRATION & PROJECT DOSSIER")
    set_run_font(tag_run, 9.5, True, GOLD)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 12.5, True, GREEN, "Georgia" if level == 1 else "Aptos")
    if level == 1:
        add_rule(p, MOSS)


def add_body(document: Document, text: str, italic: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.16
    run = p.add_run(text)
    set_run_font(run, 10.4, False, INK)
    run.italic = italic


def add_quote(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF3EC")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 10.4, True, GREEN)
    run.italic = True
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, 9.2, True, "FFFFFF")
        if widths:
            cell.width = Cm(widths[index])

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if row_index % 2 == 0:
                set_cell_shading(cell, PAPER)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            set_run_font(run, 8.9, False, INK)
            if widths:
                cell.width = Cm(widths[index])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_numbered_list(document: Document, items: list[str]) -> None:
    for number, item in enumerate(items, start=1):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{number}. {item}")
        set_run_font(run, 10.2, False, INK)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("EcoSphere AI  |  AIEM IDEAS 2026")
    set_run_font(header_run, 8.5, True, MOSS)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("EcoSphere AI — AI-Powered Sustainability Mission Control")
    set_run_font(footer_run, 8.2, False, MOSS)


def main() -> None:
    document = Document()
    configure_document(document)

    add_title(document, "EcoSphere AI", "AI-Powered Sustainability Mission Control")
    add_quote(document, "A working sustainability-operations prototype for an AIEM Campus pilot. Simulated evidence, modelled outcomes and unconfigured external services are explicitly distinguished from live, verified operations.")

    add_heading(document, "Project Snapshot")
    add_table(document, ["Item", "Details"], [
        ["Proposed Team Name", "Crimson Syndicate — please confirm this exact name before form submission."],
        ["Idea Title", "EcoSphere AI — AI-Powered Sustainability Mission Control"],
        ["Idea Category", "Artificial Intelligence, Sustainability and Smart Campus Technology"],
        ["Primary SDG", "SDG 13 — Climate Action"],
        ["Supporting SDGs", "SDG 7, SDG 9, SDG 11 and SDG 12"],
        ["Primary Beneficiaries", "Campus administrators, facilities teams, maintenance teams, sustainability clubs and students"],
        ["Prototype Positioning", "A traceable sustainability-operations prototype for an AIEM Campus pilot, with controlled simulated demonstration data when live metering is not available."],
    ], [4.0, 13.5])

    add_heading(document, "Problem Statement")
    add_body(document, "Educational campuses and institutions often receive energy, water and waste information through disconnected spreadsheets, manual records or delayed utility bills. This fragmented approach makes unusual resource consumption difficult to identify early, limits accountability for sustainability actions and makes it hard for decision-makers to compare which intervention will create the highest environmental impact for the available budget.")
    add_body(document, "The result is reactive sustainability management: a potentially avoidable HVAC spike, water leak or waste increase may be noticed only after it has already increased operational cost and carbon emissions. Institutions need a practical system that connects measurement, anomaly awareness, transparent calculations, intervention planning and evidence-led action in one understandable workflow.")

    add_heading(document, "Proposed Solution")
    add_body(document, "EcoSphere AI is an AI-powered sustainability mission-control platform designed for a campus-pilot setting. It brings energy, water and waste readings into one governed operational workflow and turns them into traceable indicators, anomaly evidence, carbon calculations, EcoScore snapshots, forecasts and clearly labelled sustainability scenarios.")
    add_quote(document, "MONITOR → DETECT → PREDICT → SIMULATE → RECOMMEND → ACT → MEASURE → REPEAT")
    add_body(document, "The product’s numerical calculations are deterministic and server-owned. AI is used to explain evidence and organise recommendations; it does not invent consumption, savings, carbon or score values. This design supports more credible sustainability discussions because every recommendation is connected to recorded or explicitly simulated evidence.")

    add_heading(document, "Core Capabilities")
    add_table(document, ["Capability", "Campus-pilot support"], [
        ["Sustainability monitoring", "Organises energy, water and waste readings against tenant, site and meter records."],
        ["Data-quality controls", "Validates imports and readings, quarantines problematic data and preserves source lineage."],
        ["Carbon and EcoScore evidence", "Applies deterministic calculation logic to create transparent carbon and EcoScore snapshots."],
        ["Anomaly detection", "Detects unusual consumption patterns and creates evidence-linked alerts."],
        ["Forecasting and trends", "Produces short-horizon, clearly limited forecasts from available evidence."],
        ["What-if simulator", "Models energy, renewable-energy, water, waste, recycling and investment choices without claiming realised outcomes."],
        ["Intervention comparison", "Compares LED upgrades, smart HVAC controls, solar installation and water-saving systems using modelled impact, cost, savings and SDG contribution."],
        ["Controlled demonstration", "Uses explicitly labelled simulated pilot data to demonstrate monitoring cycles and an HVAC-spike anomaly safely."],
    ], [4.4, 13.1])

    add_heading(document, "Innovation and Differentiation")
    add_body(document, "EcoSphere AI is not only a dashboard. It is designed as a decision-support and evidence workflow for sustainability operations. Its difference lies in connecting three layers that are commonly separated: trusted operational data, deterministic environmental calculations and user-friendly intervention planning.")
    add_body(document, "Its monitoring architecture is browser-independent, allowing eligible readings to be processed without requiring the dashboard browser to remain open. The platform distinguishes live evidence from simulated demo data and modelled scenarios, while limiting AI to explanation and recommendation reasoning. This makes the concept suitable for an initial campus pilot and for future expansion to hostels, laboratories, offices and community infrastructure.")

    add_heading(document, "Expected Beneficiaries and Value")
    add_body(document, "Campus facility and maintenance teams can use EcoSphere AI to identify potentially unusual resource consumption earlier and document follow-up actions. Administrators can use one executive view to examine resource trends, active alerts, EcoScore evidence and sustainability priorities. Students and sustainability clubs can use the platform to support transparent climate-action discussions, campaigns and pilot projects.")
    add_body(document, "The approach is also relevant to hostels, educational institutions, offices and local public facilities that need a structured, affordable route from disconnected utility data to evidence-led resource-management decisions.")

    add_heading(document, "Sustainable Development Goal Alignment")
    add_table(document, ["SDG", "EcoSphere AI contribution"], [
        ["SDG 13 — Climate Action", "Supports lower-emission operational decisions through carbon evidence, anomaly awareness and modelled intervention analysis."],
        ["SDG 7 — Affordable and Clean Energy", "Tracks electricity-related evidence and compares energy-efficiency and renewable-energy interventions."],
        ["SDG 9 — Industry, Innovation and Infrastructure", "Demonstrates a data-driven, traceable monitoring architecture for institution-scale infrastructure decisions."],
        ["SDG 11 — Sustainable Cities and Communities", "Provides a campus model that can be adapted to institutional and community infrastructure."],
        ["SDG 12 — Responsible Consumption and Production", "Connects water, waste and resource-consumption evidence with accountable improvement actions."],
    ], [6.0, 11.5])

    add_heading(document, "AIEM Campus Demonstration Flow")
    add_numbered_list(document, [
        "Open the AIEM Campus dashboard and establish the EcoScore, monitored meters and sustainability indicators.",
        "Start the clearly labelled live simulation so controlled pilot readings progress through the monitoring pipeline.",
        "Inject a controlled HVAC energy spike.",
        "Show the anomaly, alert, EcoScore response and evidence-linked AI explanation.",
        "Open the What-If Sustainability Simulator and model a 15% energy reduction.",
        "Compare projected carbon reduction, estimated savings and intervention trade-offs.",
        "Present the highest-impact intervention recommendation and its SDG alignment.",
        "Close with the roadmap from controlled simulation to an approved, live AIEM Campus meter-data pilot.",
    ])

    add_heading(document, "Implementation Approach and Responsible Boundaries")
    add_body(document, "EcoSphere AI includes a React-based user interface, an Express/tRPC backend, tenant-scoped data handling, a MySQL/TiDB-compatible persistence model and a browser-independent monitoring worker design. CSV import, controlled simulation, alerts, scenario comparison, reports, action tracking and operational administration are represented as integrated product workflows.")
    add_body(document, "For accuracy, the team should not claim that the project currently has live AIEM telemetry, certified emissions accounting, realised savings, a live Odoo integration, autonomous external notification delivery or a production scheduler unless those capabilities are separately configured, deployed and verified.", italic=True)

    add_heading(document, "Team Details")
    add_table(document, ["Role", "Name", "Roll Number", "Class", "Department"], [
        ["Project Leader", "Shivam Gawade", "24ec25", "TE", "Computer Engineering"],
        ["Member 2", "Rahul Ravi Rathod", "24ec17", "TE", "Computer Engineering"],
        ["Member 3", "Rehan Harmalkar", "24ec18", "TE", "Computer Engineering"],
        ["Member 4", "Ashwith Shetty", "24co35", "TE", "Computer Engineering"],
        ["Member 5", "Deekshith TS", "23ec12", "BE", "Computer Engineering"],
    ], [3.1, 5.0, 3.0, 2.0, 4.4])
    add_body(document, "Form item still required: enter Shivam Gawade’s active WhatsApp number in the Project Leader contact field. Verify all spellings, roll numbers, class and department format against college records before submitting.", italic=True)

    add_heading(document, "Form-Ready Responses")
    add_table(document, ["Form field", "Paste this response"], [
        ["Team Name", "Crimson Syndicate (confirm with the team before submission)"],
        ["Title of the Idea", "EcoSphere AI — AI-Powered Sustainability Mission Control"],
        ["Idea Category", "Artificial Intelligence, Sustainability and Smart Campus Technology"],
        ["Problem Statement", "Campuses and institutions often record energy, water and waste information in disconnected spreadsheets or only after a bill is received. This makes abnormal consumption difficult to identify early, limits accountability for sustainability actions, and makes it hard for decision-makers to compare the likely impact of interventions such as efficient lighting, smart HVAC controls or solar adoption. EcoSphere AI addresses this gap through a single sustainability operations platform that converts structured meter readings into traceable environmental metrics, anomaly alerts, EcoScore evidence, short-horizon forecasts and clearly labelled what-if scenarios. It helps campus teams move from delayed reporting to evidence-led monitoring, prioritisation and action."],
        ["Beneficiaries", "Educational institutions and campus facility teams are the primary beneficiaries, especially administrators, maintenance teams, sustainability clubs and students who need a shared view of resource consumption. The approach can also benefit hostels, offices, municipal facilities and small institutions that need an affordable way to monitor energy, water, waste and carbon evidence, identify unusual consumption, compare sustainability interventions and communicate progress responsibly."],
        ["SDGs to select", "Affordable and Clean Energy; Industry Innovation and Infrastructure; Sustainable Cities and Communities; Responsible Consumption and Production; Climate Action"],
    ], [4.0, 13.5])

    add_heading(document, "Registration Resources")
    add_table(document, ["Resource", "Link and intended use"], [
        ["Participants WhatsApp group", "https://chat.whatsapp.com/Bxtyka24YS8FI76RHQUIAm?mode=gi_t\nJoin only through this organiser-supplied link."],
        ["Official PPT template", "https://docs.google.com/presentation/d/1K5RmKudjFshvv-W4B9QFnShZVGWmGo_x/edit?usp=sharing&ouid=107722541612996630562&rtpof=true&sd=true\nDownload a personal copy; do not edit or share the original template."],
    ], [4.0, 13.5])

    add_heading(document, "Pre-Submission Checklist")
    add_numbered_list(document, [
        "Confirm the team name, team roster, roll numbers and department labels.",
        "Add the Project Leader’s active WhatsApp number.",
        "Select only the five SDGs listed in this dossier.",
        "Use the form-ready project text above, retaining the campus-pilot prototype framing.",
        "Check the email account shown in the form before submission.",
        "Submit only after every team member has verified their personal details.",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
